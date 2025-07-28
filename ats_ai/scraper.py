import asyncio
import os
import re
from playwright.async_api import async_playwright
from pathlib import Path
from docx import Document
from docx.shared import Inches
import datetime


class CalfusJobScraper:
    def __init__(self):
        self.base_url = "https://www.calfus.com"
        self.job_openings_url = "https://www.calfus.com/job-openings"
        self.jd_folder = Path("jd_folder")

    async def setup_folder(self):
        """Create jd_folder if it doesn't exist"""
        self.jd_folder.mkdir(exist_ok=True)
        print(f"Created/verified folder: {self.jd_folder}")

    def sanitize_filename(self, title):
        """Sanitize job title for filename"""
        # Remove special characters and replace spaces with underscores
        sanitized = re.sub(r'[<>:"/\\|?*]', '', title)
        sanitized = re.sub(r'\s+', '_', sanitized.strip())
        # Remove common unwanted prefixes
        sanitized = sanitized.replace('Current_Job_Openings_', '')
        return sanitized

    def clean_job_content(self, content):
        """Clean and filter job content to remove unwanted sections"""
        if not content:
            return ""

        lines = content.split('\n')
        cleaned_lines = []
        skip_current_section = False

        # Keywords that indicate sections to skip
        skip_section_keywords = [
            'about us', 'about the company', 'company overview', 'who we are',
            'our company', 'company profile', 'organization overview',
            'apply now', 'how to apply', 'application process',
            'source url', 'scraped date', 'contact us', 'get in touch',
            'follow us', 'social media', 'connect with us',
            'navigation', 'menu', 'home', 'careers', 'lets connect',
            'join us', 'our work', 'agent foundry', 'benefits and perks'
        ]

        # Patterns that indicate metadata or unwanted content
        skip_patterns = [
            r'^\*\*Source URL:\s*\*\*',
            r'^\*\*Scraped Date:\s*\*\*',
            r'^https?://',
            r'^\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}:\d{2}',
            r'^Apply Now\s*$',
            r'^Back\s*$',
            r'^Home\s*$',
            r'^Menu\s*$',
            r'^Navigation\s*$'
        ]

        for line in lines:
            line = line.strip()

            # Skip empty lines initially, we'll add them back strategically
            if not line:
                if cleaned_lines and cleaned_lines[-1]:  # Only add empty line if previous line wasn't empty
                    cleaned_lines.append('')
                continue

            # Check if line matches skip patterns
            should_skip_line = False
            for pattern in skip_patterns:
                if re.match(pattern, line, re.IGNORECASE):
                    should_skip_line = True
                    break

            if should_skip_line:
                continue

            # Check if this line starts a section we want to skip
            line_lower = line.lower()

            # Reset skip flag for new sections
            if line.endswith(':') or re.match(r'^[A-Z][^a-z]*:?$', line):
                skip_current_section = any(keyword in line_lower for keyword in skip_section_keywords)

            # Skip if we're in a section to skip
            if skip_current_section:
                # Check if we've moved to a new section that we want to keep
                if line.endswith(':') and not any(keyword in line_lower for keyword in skip_section_keywords):
                    skip_current_section = False
                    cleaned_lines.append(line)
                continue

            # Skip lines that are just navigation or metadata
            if any(keyword == line_lower for keyword in skip_section_keywords):
                continue

            # Skip very short lines that are likely navigation
            if len(line) <= 2 and line.isalpha():
                continue

            # Skip lines that are just dates or URLs
            if re.match(r'^\d{4}-\d{2}-\d{2}', line) or line.startswith('http'):
                continue

            cleaned_lines.append(line)

        # Remove multiple consecutive empty lines
        final_lines = []
        prev_empty = False

        for line in cleaned_lines:
            if not line.strip():
                if not prev_empty:
                    final_lines.append(line)
                prev_empty = True
            else:
                final_lines.append(line)
                prev_empty = False

        # Remove leading and trailing empty lines
        while final_lines and not final_lines[0].strip():
            final_lines.pop(0)
        while final_lines and not final_lines[-1].strip():
            final_lines.pop()

        return '\n'.join(final_lines)

    async def scrape_job_listings(self, page):
        """Scrape all job listings from the main page"""
        max_retries = 3
        retry_delay = 5

        for attempt in range(max_retries):
            try:
                print(f"Attempting to load job openings page (attempt {attempt + 1}/{max_retries})")
                await page.goto(self.job_openings_url, wait_until="networkidle", timeout=60000)
                await page.wait_for_timeout(5000)  # Increased wait time
                print("Successfully loaded job openings page")
                break
            except Exception as e:
                print(f"Attempt {attempt + 1} failed: {e}")
                if attempt < max_retries - 1:
                    print(f"Retrying in {retry_delay} seconds...")
                    await asyncio.sleep(retry_delay)
                    retry_delay *= 2  # Exponential backoff
                else:
                    print("All attempts failed. Please check your internet connection and try again.")
                    raise e

        # Check if page has loaded properly and scroll to load any lazy content
        await page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
        await page.wait_for_timeout(3000)
        await page.evaluate("window.scrollTo(0, 0)")
        await page.wait_for_timeout(2000)

        jobs = []

        # Try multiple approaches to find job listings
        print("Analyzing page structure...")

        # Method 1: Look for "See Details" text in various elements
        see_details_selectors = [
            'a:has-text("See Details")',
            'button:has-text("See Details")',
            'a:has-text("see details")',
            'button:has-text("see details")',
            'a[href*="job"]',
            'a[href*="opening"]',
            'a[href*="position"]',
            '*:has-text("See Details")',
            '*:has-text("see details")',
            '*:has-text("Apply")',
            '*:has-text("View")'
        ]

        all_job_elements = []
        for selector in see_details_selectors:
            try:
                elements = await page.query_selector_all(selector)
                print(f"Selector '{selector}' found {len(elements)} elements")
                for element in elements:
                    if element not in all_job_elements:
                        all_job_elements.append(element)
            except Exception as e:
                print(f"Error with selector '{selector}': {e}")
                continue

        print(f"Total unique job-related elements found: {len(all_job_elements)}")

        # Method 2: Look for job containers by structure
        print("Looking for job containers by structure...")
        job_containers = await page.query_selector_all('div, section, article')
        potential_job_containers = []

        for container in job_containers:
            try:
                container_text = await container.inner_text()
                container_text_lower = container_text.lower()

                # Check if container has job-like content
                job_indicators = [
                    'see details', 'apply now', 'view job', 'job opening',
                    'position', 'bengaluru', 'pune', 'mumbai', 'hyderabad',
                    'years of experience', 'experience', 'skills required'
                ]

                if any(indicator in container_text_lower for indicator in job_indicators):
                    # Check if it's not a navigation or header element
                    if (not any(nav in container_text_lower for nav in ['navigation', 'menu', 'header', 'footer'])
                            and len(container_text.strip()) > 50):
                        potential_job_containers.append(container)
            except:
                continue

        print(f"Found {len(potential_job_containers)} potential job containers")

        # Process all found elements
        processed_urls = set()  # To avoid duplicates

        for i, element in enumerate(all_job_elements + potential_job_containers):
            try:
                job_title = ""
                location = ""
                job_url = ""

                # Try to get URL from the element itself or find a link within it
                if await element.evaluate('el => el.tagName.toLowerCase()') == 'a':
                    job_url = await element.get_attribute('href') or ""
                else:
                    # Look for links within the element
                    link_element = await element.query_selector('a[href]')
                    if link_element:
                        job_url = await link_element.get_attribute('href') or ""

                # Skip if no URL found or if it's a duplicate
                if not job_url or job_url in processed_urls:
                    continue

                # Make URL absolute
                if job_url.startswith('/'):
                    job_url = self.base_url + job_url

                # Skip if it doesn't look like a job URL
                if not any(keyword in job_url.lower() for keyword in ['job', 'opening', 'position', 'career']):
                    continue

                processed_urls.add(job_url)

                # Find job title and location by traversing up the DOM
                current_element = element
                for _ in range(10):  # Try up to 10 parent levels
                    try:
                        element_text = await current_element.inner_text()
                        lines = [line.strip() for line in element_text.split('\n') if line.strip()]

                        # Look for job title (usually the first meaningful line)
                        for line in lines:
                            if (len(line) > 5 and len(line) < 100 and
                                    not line.lower() in ['see details', 'apply now', 'view job'] and
                                    not any(keyword in line.lower() for keyword in
                                            ['current job openings', 'calfus', 'navigation', 'menu'])):
                                if not job_title:
                                    job_title = line
                                    break

                        # Look for location
                        for line in lines:
                            line_lower = line.lower()
                            if 'bengaluru' in line_lower or 'bangalore' in line_lower:
                                location = "Bengaluru"
                                break
                            elif 'pune' in line_lower:
                                location = "Pune"
                                break
                            elif 'mumbai' in line_lower:
                                location = "Mumbai"
                                break
                            elif 'hyderabad' in line_lower:
                                location = "Hyderabad"
                                break

                        if job_title and location:
                            break

                        # Move to parent element
                        parent = await current_element.query_selector('..')
                        if parent:
                            current_element = parent
                        else:
                            break
                    except:
                        break

                # If still no title found, use a fallback
                if not job_title:
                    job_title = f"Job_Position_{len(jobs) + 1}"

                if not location:
                    location = "Location_TBD"

                jobs.append({
                    'title': job_title,
                    'location': location,
                    'url': job_url
                })
                print(f"Found job {len(jobs)}: {job_title} in {location}")

            except Exception as e:
                print(f"Error processing element {i}: {e}")
                continue

        # Remove duplicates based on URL
        seen_urls = set()
        unique_jobs = []
        for job in jobs:
            if job['url'] not in seen_urls:
                seen_urls.add(job['url'])
                unique_jobs.append(job)

        print(f"Final count: {len(unique_jobs)} unique jobs found")

        # If still no jobs found, take a screenshot for debugging
        if not unique_jobs:
            print("No jobs found. Taking screenshot and saving page content for debugging...")
            await page.screenshot(path="debug_page.png", full_page=True)
            page_content = await page.content()
            with open("debug_page.html", "w", encoding="utf-8") as f:
                f.write(page_content)
            print("Debug files saved: debug_page.png and debug_page.html")

        return unique_jobs

    async def scrape_job_details(self, page, job):
        """Scrape individual job description"""
        max_retries = 3
        retry_delay = 3

        for attempt in range(max_retries):
            try:
                print(f"Scraping details for: {job['title']} (attempt {attempt + 1}/{max_retries})")
                await page.goto(job['url'], wait_until="networkidle", timeout=60000)
                await page.wait_for_timeout(3000)
                break
            except Exception as e:
                print(f"Attempt {attempt + 1} failed for {job['title']}: {e}")
                if attempt < max_retries - 1:
                    print(f"Retrying in {retry_delay} seconds...")
                    await asyncio.sleep(retry_delay)
                    retry_delay *= 2
                else:
                    print(f"All attempts failed for {job['title']}")
                    return {
                        'title': job['title'],
                        'location': job['location'],
                        'url': job['url'],
                        'page_title': 'Error - Connection Failed',
                        'content': f"Error: Could not connect to job page after {max_retries} attempts",
                        'scraped_date': datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    }

        try:
            # Get the page title for better job identification
            page_title = await page.title()

            # Try to get the main content
            content_selectors = [
                'main',
                '.main-content',
                '.content',
                '.job-description',
                '.job-details',
                'article',
                '.container'
            ]

            job_content = ""
            for selector in content_selectors:
                element = await page.query_selector(selector)
                if element:
                    content = await element.inner_text()
                    if content and len(content.strip()) > 200:
                        job_content = content.strip()
                        break

            # If no good content found, get body text and filter
            if not job_content or len(job_content) < 200:
                body_text = await page.evaluate('document.body.innerText')
                job_content = body_text

            # Clean the content to remove unwanted sections
            cleaned_content = self.clean_job_content(job_content)

            return {
                'title': job['title'],
                'location': job['location'],
                'url': job['url'],
                'page_title': page_title,
                'content': cleaned_content,
                'scraped_date': datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }

        except Exception as e:
            print(f"Error scraping job details for {job['title']}: {e}")
            return {
                'title': job['title'],
                'location': job['location'],
                'url': job['url'],
                'page_title': 'Error',
                'content': f"Error scraping job: {str(e)}",
                'scraped_date': datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }

    def save_job_as_docx(self, job_data):
        """Save job description as Word document"""
        try:
            # Create a new Document
            doc = Document()

            # Add title
            title = doc.add_heading(job_data['title'], 0)

            # Add location
            location_para = doc.add_paragraph()
            location_para.add_run('Location: ').bold = True
            location_para.add_run(f"{job_data['location']}")

            # Add job description (cleaned content)
            doc.add_heading('Job Description', level=1)

            # Split content into paragraphs and add them
            if job_data['content']:
                content_paragraphs = job_data['content'].split('\n')
                current_paragraph = ""

                for line in content_paragraphs:
                    if line.strip():
                        if current_paragraph:
                            current_paragraph += " " + line.strip()
                        else:
                            current_paragraph = line.strip()
                    else:
                        # Empty line indicates paragraph break
                        if current_paragraph:
                            doc.add_paragraph(current_paragraph)
                            current_paragraph = ""

                # Add the last paragraph if it exists
                if current_paragraph:
                    doc.add_paragraph(current_paragraph)

            # Save the document
            filename = f"{self.sanitize_filename(job_data['title'])}.docx"
            filepath = self.jd_folder / filename

            doc.save(str(filepath))
            print(f"Saved: {filepath}")

        except Exception as e:
            print(f"Error saving document for {job_data['title']}: {e}")

    async def run(self):
        """Main scraping function"""
        await self.setup_folder()

        async with async_playwright() as p:
            # Launch browser with additional options
            browser = await p.chromium.launch(
                headless=False,
                args=[
                    '--no-sandbox',
                    '--disable-dev-shm-usage',
                    '--disable-web-security',
                    '--disable-features=VizDisplayCompositor'
                ]
            )
            context = await browser.new_context(
                user_agent="Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                viewport={'width': 1920, 'height': 1080},
                extra_http_headers={
                    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
                    'Accept-Language': 'en-US,en;q=0.5',
                    'Accept-Encoding': 'gzip, deflate',
                    'DNT': '1',
                    'Connection': 'keep-alive',
                    'Upgrade-Insecure-Requests': '1'
                }
            )
            page = await context.new_page()

            try:
                # Test basic connectivity first
                print("Testing connectivity to calfus.com...")
                try:
                    await page.goto("https://www.calfus.com", timeout=30000)
                    print("Successfully connected to calfus.com")
                    await page.wait_for_timeout(2000)
                except Exception as e:
                    print(f"Failed to connect to main site: {e}")
                    print("This might be a network connectivity issue or the site might be down.")
                    return

                # Scrape job listings
                print("\nScraping job listings...")
                jobs = await self.scrape_job_listings(page)

                if not jobs:
                    print("No jobs found. Taking screenshot for debugging...")
                    await page.screenshot(path="debug_page.png")
                    print("Please check debug_page.png to see what the page looks like")
                    return

                print(f"Found {len(jobs)} jobs to scrape")

                # Scrape each job's details
                successful_scrapes = 0
                for i, job in enumerate(jobs, 1):
                    print(f"\nProcessing job {i}/{len(jobs)}")
                    job_data = await self.scrape_job_details(page, job)

                    # Only save if we got meaningful content
                    if job_data['content'] and len(job_data['content'].strip()) > 50:
                        self.save_job_as_docx(job_data)
                        successful_scrapes += 1
                    else:
                        print(f"Warning: No meaningful content found for {job['title']}")

                    # Add delay between requests
                    await asyncio.sleep(3)

                print(
                    f"\nCompleted! Successfully scraped {successful_scrapes} out of {len(jobs)} job descriptions to {self.jd_folder}")

            except Exception as e:
                print(f"Error during scraping: {e}")
                import traceback
                traceback.print_exc()

            finally:
                await browser.close()


async def main():
    scraper = CalfusJobScraper()
    await scraper.run()


if __name__ == "__main__":
    # Install required packages first:
    # pip install playwright python-docx
    # playwright install

    asyncio.run(main())